"""
Document processing utilities for extracting text from various file formats.
"""
import os
import mimetypes
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
import asyncio
from dataclasses import dataclass

# Text extraction libraries
import PyPDF2
from docx import Document
import markdownify
from bs4 import BeautifulSoup

from .chunking import TextChunker, Chunk


@dataclass
class ProcessedDocument:
    """Represents a processed document with extracted content."""
    
    content: str
    metadata: Dict[str, Any]
    file_path: str
    file_type: str
    chunks: Optional[List[Chunk]] = None


class DocumentProcessor:
    """Processor for extracting text from various document formats."""
    
    SUPPORTED_FORMATS = {
        '.pdf': 'application/pdf',
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.html': 'text/html',
        '.htm': 'text/html',
    }
    
    def __init__(self, chunker: Optional[TextChunker] = None):
        """
        Initialize the document processor.
        
        Args:
            chunker: Text chunker instance for splitting documents
        """
        self.chunker = chunker or TextChunker()
    
    def is_supported(self, file_path: str) -> bool:
        """
        Check if a file format is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if supported, False otherwise
        """
        extension = Path(file_path).suffix.lower()
        return extension in self.SUPPORTED_FORMATS
    
    async def process_file(self, file_path: str) -> ProcessedDocument:
        """
        Process a single file and extract its content.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            ProcessedDocument object
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not self.is_supported(file_path):
            raise ValueError(f"Unsupported file format: {Path(file_path).suffix}")
        
        # Extract basic metadata
        file_stat = os.stat(file_path)
        metadata = {
            "file_name": Path(file_path).name,
            "file_path": file_path,
            "file_size": file_stat.st_size,
            "modified_time": file_stat.st_mtime,
            "file_extension": Path(file_path).suffix.lower(),
        }
        
        # Extract content based on file type
        extension = Path(file_path).suffix.lower()
        
        if extension == '.pdf':
            content = await self._extract_pdf(file_path)
        elif extension == '.docx':
            content = await self._extract_docx(file_path)
        elif extension in ['.html', '.htm']:
            content = await self._extract_html(file_path)
        elif extension == '.md':
            content = await self._extract_markdown(file_path)
        else:  # .txt and other text files
            content = await self._extract_text(file_path)
        
        # Add content-based metadata
        metadata.update({
            "content_length": len(content),
            "word_count": len(content.split()) if content else 0,
        })
        
        return ProcessedDocument(
            content=content,
            metadata=metadata,
            file_path=file_path,
            file_type=extension
        )
    
    async def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        def extract():
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    
                    return '\n\n'.join(text_parts)
            except Exception as e:
                raise Exception(f"Failed to extract PDF content: {str(e)}")
        
        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)
    
    async def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        def extract():
            try:
                doc = Document(file_path)
                text_parts = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(' | '.join(row_text))
                
                return '\n\n'.join(text_parts)
            except Exception as e:
                raise Exception(f"Failed to extract DOCX content: {str(e)}")
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)
    
    async def _extract_html(self, file_path: str) -> str:
        """Extract text from HTML file."""
        def extract():
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    html_content = file.read()
                
                # Parse HTML and extract text
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Get text and clean it up
                text = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                return text
            except Exception as e:
                raise Exception(f"Failed to extract HTML content: {str(e)}")
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)
    
    async def _extract_markdown(self, file_path: str) -> str:
        """Extract text from Markdown file."""
        def extract():
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    markdown_content = file.read()
                
                # Convert markdown to HTML then to plain text
                html = markdownify.markdownify(markdown_content)
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text()
            except Exception as e:
                # Fallback to reading as plain text
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)
    
    async def _extract_text(self, file_path: str) -> str:
        """Extract text from plain text file."""
        def extract():
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            except UnicodeDecodeError:
                # Try with different encoding
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                raise Exception(f"Failed to extract text content: {str(e)}")
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)
    
    async def process_directory(
        self,
        directory_path: str,
        recursive: bool = True,
        max_files: Optional[int] = None
    ) -> List[ProcessedDocument]:
        """
        Process all supported files in a directory.
        
        Args:
            directory_path: Path to the directory
            recursive: Whether to process subdirectories
            max_files: Maximum number of files to process
            
        Returns:
            List of ProcessedDocument objects
        """
        if not os.path.exists(directory_path):
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        files_to_process = []
        
        if recursive:
            for root, dirs, files in os.walk(directory_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    if self.is_supported(file_path):
                        files_to_process.append(file_path)
        else:
            for item in os.listdir(directory_path):
                item_path = os.path.join(directory_path, item)
                if os.path.isfile(item_path) and self.is_supported(item_path):
                    files_to_process.append(item_path)
        
        # Limit number of files if specified
        if max_files:
            files_to_process = files_to_process[:max_files]
        
        # Process files concurrently
        processed_documents = []
        semaphore = asyncio.Semaphore(5)  # Limit concurrent processing
        
        async def process_with_semaphore(file_path):
            async with semaphore:
                try:
                    return await self.process_file(file_path)
                except Exception as e:
                    print(f"Error processing {file_path}: {str(e)}")
                    return None
        
        tasks = [process_with_semaphore(file_path) for file_path in files_to_process]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filter out None results and exceptions
        for result in results:
            if isinstance(result, ProcessedDocument):
                processed_documents.append(result)
        
        return processed_documents
    
    async def chunk_document(
        self,
        processed_doc: ProcessedDocument,
        document_id: Optional[str] = None
    ) -> ProcessedDocument:
        """
        Chunk a processed document.
        
        Args:
            processed_doc: ProcessedDocument to chunk
            document_id: Custom document ID
            
        Returns:
            ProcessedDocument with chunks added
        """
        if not processed_doc.content:
            processed_doc.chunks = []
            return processed_doc
        
        doc_id = document_id or processed_doc.metadata.get("file_name", "unknown")
        
        chunks = self.chunker.chunk_text(
            text=processed_doc.content,
            metadata=processed_doc.metadata,
            document_id=doc_id
        )
        
        processed_doc.chunks = chunks
        return processed_doc
    
    async def process_and_chunk_file(
        self,
        file_path: str,
        document_id: Optional[str] = None
    ) -> ProcessedDocument:
        """
        Process and chunk a file in one step.
        
        Args:
            file_path: Path to the file
            document_id: Custom document ID
            
        Returns:
            ProcessedDocument with chunks
        """
        processed_doc = await self.process_file(file_path)
        return await self.chunk_document(processed_doc, document_id)
    
    def get_processing_stats(self, documents: List[ProcessedDocument]) -> Dict[str, Any]:
        """
        Get statistics about processed documents.
        
        Args:
            documents: List of processed documents
            
        Returns:
            Dictionary with processing statistics
        """
        if not documents:
            return {}
        
        total_files = len(documents)
        total_content_length = sum(len(doc.content) for doc in documents)
        total_word_count = sum(doc.metadata.get("word_count", 0) for doc in documents)
        
        file_types = {}
        for doc in documents:
            file_type = doc.file_type
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        # Chunk statistics
        total_chunks = 0
        for doc in documents:
            if doc.chunks:
                total_chunks += len(doc.chunks)
        
        return {
            "total_files": total_files,
            "total_content_length": total_content_length,
            "total_word_count": total_word_count,
            "avg_content_length": total_content_length / total_files if total_files > 0 else 0,
            "avg_word_count": total_word_count / total_files if total_files > 0 else 0,
            "file_types": file_types,
            "total_chunks": total_chunks,
            "avg_chunks_per_document": total_chunks / total_files if total_files > 0 else 0,
        }